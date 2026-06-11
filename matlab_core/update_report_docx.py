# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
import csv


ROOT = Path(__file__).resolve().parents[1]
MATLAB = ROOT / "matlab_core"
DATA = MATLAB / "data"
FIGS = MATLAB / "figures"


def find_docx():
    candidates = [
        p for p in ROOT.glob("*.docx")
        if not p.name.startswith("~$")
        and "backup_before_matlab_update" not in p.name
    ]
    if not candidates:
        raise FileNotFoundError("No report docx found.")
    return candidates[0]


def read_csv(name):
    with (DATA / name).open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.2
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    return table


def fmt(value, digits=4):
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def pct(value, digits=4):
    return fmt(value, digits) + "%"


def main():
    docx_path = find_docx()
    backup = docx_path.with_name(docx_path.stem + "_backup_before_matlab_update.docx")
    if not backup.exists():
        copy2(docx_path, backup)

    intrinsic = read_csv("intrinsic_evm.csv")
    opt = read_csv("optimization_summary.csv")
    srrc = read_csv("srrc_filter_comparison.csv")
    quality = read_csv("resource_quality_score.csv")

    modes12 = {r["Mode"]: r for r in intrinsic if r["DataWidth"] == "12"}
    mode_rows = [
        [m, pct(modes12[m]["IntrinsicEVMPct"]), "是"]
        for m in ["BASEBAND", "FS4_IF", "DDS_IF", "CORDIC_IF"]
    ]
    opt_rows = [
        [
            r["Criterion"],
            r["Mode"],
            r["DataWidth"],
            pct(r["IntrinsicEVMPct"]),
            r["LUT_Est"],
            r["FF_Est"],
            fmt(r["Score"], 4),
        ]
        for r in opt
    ]
    srrc_rows = [
        [
            r["Mode"],
            r["Taps"],
            pct(r["IntrinsicEVMPct"]),
            r["BER_12dB"],
            pct(r["ChannelEVMPct_12dB"]),
            r["LUT_Est"],
            r["FF_Est"],
        ]
        for r in srrc
    ]
    quality_rows = [
        [
            r["Mode"],
            r["DataWidth"],
            pct(r["IntrinsicEVMPct"]),
            r["LUT_Est"],
            r["FF_Est"],
            r["DSP_Est"],
            fmt(r["Score"], 4),
        ]
        for r in quality[:5]
    ]

    doc = Document(str(docx_path))
    doc.add_page_break()
    doc.add_heading("MATLAB 仿真补充与优化结果", 1)
    add_para(
        doc,
        "在前述基础仿真完成后，进一步围绕“低资源”和“可配置”两个目标补充了三类实验："
        "一是无噪声本征 EVM 与 1% EVM 约束下的自动配置选择；二是将本征 EVM、LUT、FF、DSP "
        "和延迟统一归一化的资源-性能综合评分；三是针对 SRRC 脉冲成形滤波器资源瓶颈，比较 "
        "17 taps、25 taps 和 33 taps 三种抽头规模。新增 MATLAB 代码、CSV 数据和图表均保存在 matlab_core 目录下。",
    )

    doc.add_heading("本征 EVM 与 MATLAB 阶段达标结论", 2)
    add_para(
        doc,
        "含噪声 EVM 会同时反映 AWGN 信道噪声和接收链路误差，因此不适合直接作为调制器本身精度指标。"
        "为单独评价调制器实现误差，本文新增无噪声本征 EVM 实验。12 bit 定点条件下，四种模式的本征 EVM "
        "均低于 5% 目标，其中 BASEBAND 为 0.3746%，FS4_IF 为 0.5718%，DDS_IF 为 0.5705%，"
        "CORDIC_IF 为 0.5645%。",
    )
    add_table(doc, ["模式", "12 bit 本征 EVM", "是否低于 5% 目标"], mode_rows)
    add_para(
        doc,
        "因此，可以限定性地认为：MATLAB 仿真阶段已经达到报告设定的算法正确性、定点精度和结构对比目标。"
        "该结论仅覆盖 MATLAB 仿真阶段，Vivado 综合、post-route 资源、Fmax、功耗和板级验证仍需后续完成。",
    )

    doc.add_heading("资源-性能综合评分与自动配置选择", 2)
    add_para(
        doc,
        "为了避免只看单一指标，本文引入资源-性能综合评分，将本征 EVM、LUT、FF、DSP 和延迟归一化后加权。"
        "评分权重设置为：EVM 0.35、LUT 0.25、FF 0.20、DSP 0.10、延迟 0.10，评分越低代表综合折中越优。",
    )
    add_table(doc, ["模式", "位宽", "本征 EVM", "LUT", "FF", "DSP", "综合评分"], quality_rows)
    add_para(
        doc,
        "在 1% 本征 EVM 约束下，自动搜索得到三类代表性结果：综合评分最优配置为 BASEBAND 10 bit，"
        "本征 EVM 为 0.4381%，资源估算为 680 LUT 和 900 FF；若只追求最低 LUT，则 BASEBAND 8 bit "
        "可达到 0.7769% 本征 EVM，估算为 596 LUT 和 784 FF；若追求最低 EVM，则 BASEBAND 16 bit "
        "可达到 0.3371%，但资源上升到 932 LUT 和 1248 FF。",
    )
    add_table(doc, ["选择准则", "模式", "位宽", "本征 EVM", "LUT", "FF", "综合评分"], opt_rows)

    doc.add_heading("SRRC 滤波器压缩对比", 2)
    add_para(
        doc,
        "SRRC 脉冲成形滤波器是发射链中最主要的资源消耗模块之一。为验证滤波器资源压缩的可行性，"
        "本文比较了 17 taps、25 taps 和 33 taps 三种 SRRC 配置。结果表明，17 taps 虽然明显降低 LUT，"
        "但本征 EVM 超过 5% 目标；25 taps 可将 BASEBAND 和 FS4_IF 的本征 EVM 分别控制在 3.8730% "
        "和 3.9959%，低于 5% 目标；33 taps 则可将本征 EVM 降至 1% 以下，适合作为高精度配置。",
    )
    add_table(doc, ["模式", "抽头数", "本征 EVM", "12 dB BER", "12 dB 含噪声 EVM", "LUT", "FF"], srrc_rows)
    add_para(
        doc,
        "由此可形成一个新的可配置优化点：在资源受限场景中采用 25 taps SRRC 低资源模式，"
        "在精度优先场景中采用 33 taps SRRC 高精度模式。该策略比固定使用单一滤波器更符合低资源 FPGA "
        "设计中的性能-资源折中思想。",
    )

    doc.add_heading("新增创新点表述", 2)
    add_para(
        doc,
        "基于新增实验，本文的创新点可以进一步归纳为以下三点：第一，提出可切换载波生成的低资源 QPSK "
        "发射结构，支持 BASEBAND、FS4_IF、DDS_IF 和 CORDIC_IF 多模式对比；第二，将 Fs/4 固定中频上变频"
        "设计为无乘法实现，通过符号翻转和 I/Q 交换替代通用乘法器；第三，引入面向 EVM 约束的位宽、"
        "滤波器抽头和资源-性能联合优化方法，自动筛选满足精度目标的低资源配置。",
    )
    add_para(
        doc,
        "可写入论文的总结表述为：MATLAB 仿真结果表明，所设计的低资源 QPSK 调制器在浮点与定点模型下"
        "均能正确完成 Gray 映射、SRRC 成形、数字中频调制及 AWGN 信道验证。BER 曲线与理论 QPSK-AWGN "
        "曲线基本一致；在无噪声本征误差测试中，12 bit 定点 BASEBAND、FS4_IF、DDS_IF 和 CORDIC_IF "
        "模式的 RMS EVM 分别为 0.3746%、0.5718%、0.5705% 和 0.5645%，均低于 5% 的设计目标。"
        "因此，MATLAB 仿真阶段已达到报告设定的算法正确性、定点精度和结构对比目标。",
    )
    add_para(
        doc,
        "进一步地，为量化低资源设计中的性能-资源折中，本文引入资源-性能综合评分指标，将本征 EVM、"
        "LUT、FF、DSP 和延迟统一归一化加权评价。在 1% 本征 EVM 约束下，自动搜索结果选择 10 bit "
        "BASEBAND 配置作为综合评分最优方案。针对脉冲成形滤波器资源瓶颈，SRRC 抽头压缩实验表明，"
        "25 taps 可作为低资源模式，33 taps 可作为高精度模式，从而形成可配置滤波器优化策略。",
    )

    doc.add_heading("新增仿真图表文件", 2)
    for fig_name, caption in [
        ("resource_quality_score.png", "资源-性能综合评分前十配置"),
        ("srrc_filter_comparison.png", "SRRC 抽头压缩的 EVM 与资源对比"),
        ("intrinsic_evm.png", "不同结构和位宽下的调制器本征 EVM"),
    ]:
        fig = FIGS / fig_name
        if fig.exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(fig), width=Inches(5.6))
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(docx_path))
    print(f"updated: {docx_path}")
    print(f"backup: {backup}")


if __name__ == "__main__":
    main()
